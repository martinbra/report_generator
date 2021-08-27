# MIT License
#
# Copyright (c) 2021 Martin Meier
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.

# Install needed libraries:
# pip install python-docx PIL glob
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageTk
import glob
import tkinter as tk
from tkinter import filedialog, ttk
from random import choice
from string import ascii_uppercase
import os


def selecionar_pasta():
    """ Permite ao usuário selecionar uma pasta para varredura, retorna o caminho da pasta. """
    root = tk.Tk()
    root.withdraw()
    root.folder_path = filedialog.askdirectory()
    root.destroy()
    return root.folder_path


def gerar_documento(pasta_escolhida, nome_do_documento, lista_de_fotos):
    """ Gera um documento DOCX com nome e fotos providas, salva na pasta das fotos """

    documento = Document()

    documento.add_heading(nome_do_documento, 0)

    p = documento.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    documento.add_heading('Heading, level 1', level=1)
    documento.add_paragraph('Intense quote', style='Intense Quote')

    documento.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    documento.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    # Adiciona Imagens
    for foto in lista_de_fotos:
        documento.add_picture(foto, width=Inches(5.5))

    # Salva documento
    caminho_do_arquivo = os.path.join(pasta_escolhida, f"{nome_do_documento}.docx")
    documento.save(caminho_do_arquivo)


def detecta_foto_preta(foto, tolerancia=0.2):
    """ Calcula se uma foto é "preta" com tolerancia entre 0 e 1.0 """
    img = Image.open(foto)
    limite = round(tolerancia * 255)

    # Função de filtro para detecção de preto/branco.
    def filtro(x):
        if x > limite:
            return 255
        else:
            return 0

    # Converte uma imagem para preto e branco utilizando função de filtro.
    img_preto_e_branco = img.convert('L').point(filtro, mode='1')

    # Há algum pixel branco?
    tem_branco = any(img_preto_e_branco.getdata())

    return not tem_branco


def agrupar_fotos(lista_de_fotos):
    """ Agrupo fotos da lista recebida, separando grupos ao detectar uma foto preta. """
    novo_grupo = []
    grupos_de_fotos = [novo_grupo]

    for foto in lista_de_fotos:
        if detecta_foto_preta(foto):
            # Foto preta divide grupo de fotos.
            # Se grupo de fotos conter alguma foto válida, inicia novo grupo.
            if len(novo_grupo) > 0:
                novo_grupo = []
                grupos_de_fotos.append(novo_grupo)
        else:
            # Foto válida, adiciona ao grupo.
            novo_grupo.append(foto)

    return grupos_de_fotos


def mostrar_fotos_obter_nome(grupo_de_fotos):
    """ Mostra fotos do grupo em uma aba para obter nome do arquivo. """

    # TODO: remover esse global. Depende de melhorias na função guarda_nome_do_arquivo
    global nome_para_arquivo
    nome_para_arquivo = ""

    janela = tk.Tk()
    janela.geometry("1024x768")
    # janela.state('zoomed')

    def imagem_maximizada(foto):
        janela_x = 1024  # janela.winfo_width()
        janela_y = 768  # janela.winfo_height()
        # print(janela_x, janela_y)

        img = Image.open(foto)
        img_x, img_y = img.size
        # print(img_x, img_y)
        escala = min(janela_x/img_x, janela_y/img_y) * 0.8
        img_x, img_y = round(img_x * escala), round(img_y * escala)
        # print(img_x, img_y)

        img = img.resize((img_x, img_y))

        return ImageTk.PhotoImage(img)

    def guarda_nome_do_arquivo(*args):
        _ = args  # Unused Tkinter arguments
        # TODO: remover esse global. Depende de melhorias na função guarda_nome_do_arquivo
        global nome_para_arquivo
        nome_para_arquivo = nome.get()
        janela.destroy()

    # TODO:
    #  ativar botão OK (e também ENTER)
    #  escala das imagens dinâmica
    #  https://stackoverflow.com/questions/24061099/tkinter-resize-background-image-to-window-size

    # get value
    frame = tk.Frame(janela)
    frame.pack(side=tk.TOP, fill=tk.X)
    ttk.Label(frame, text="Nome do arquivo a gerar:").pack(side=tk.LEFT)
    nome = ttk.Entry(frame, textvariable=nome_para_arquivo)
    nome.pack(side=tk.LEFT, fill=tk.X, expand=True)
    nome.bind('<Return>', guarda_nome_do_arquivo)
    nome.focus_set()
    botao = ttk.Button(frame, text='OK', command=guarda_nome_do_arquivo)
    botao.pack(side=tk.LEFT)

    controle_abas = ttk.Notebook(janela)

    # Gera objetos das imagens para usar nas abas.
    # Caso imagem fosse gerada dentro do loop, apenas a foto da última aba apareceria.
    fotos = [imagem_maximizada(foto) for foto in grupo_de_fotos]

    for idx, (foto, caminho_da_foto) in enumerate(zip(fotos, grupo_de_fotos)):
        nova_aba = ttk.Frame(controle_abas)
        ttk.Label(nova_aba, text=caminho_da_foto).pack()
        ttk.Label(nova_aba, image=foto).pack()
        controle_abas.add(nova_aba, text=f"{idx:02}")

    controle_abas.pack(expand=1, fill="both")

    # Trás janela para frente, e foca nela.
    janela.after(1, lambda: janela.focus_force())

    janela.protocol("WM_DELETE_WINDOW", guarda_nome_do_arquivo)
    janela.mainloop()

    # Limpa nome do arquivo
    nome_para_arquivo = "".join(c for c in nome_para_arquivo if c.isalnum() or c in ' ._').strip()

    if not nome_para_arquivo:
        nome_para_arquivo = "arquivo_sem_nome"

    # TODO: retornar nome fornecido por usuário em campo de texto.
    return nome_para_arquivo


def move_arquivos(pasta_escolhida, nome_do_arquivo, grupo_de_fotos):

    # Cria nova pasta
    nova_pasta = os.path.join(pasta_escolhida, nome_do_arquivo)
    os.mkdir(nova_pasta)

    # Move Fotos
    for foto in grupo_de_fotos:
        novo_caminho = os.path.join(nova_pasta, os.path.basename(foto))
        os.replace(foto, novo_caminho)


def rodar_programa():

    # TODO: Splashscreen, menu de inicialização.

    # Permite ao usuário selecionar uma pasta, e busca fotos da pasta.
    pasta_escolhida = selecionar_pasta()
    fotos_da_pasta = glob.glob(f"{pasta_escolhida}/*.jpg")

    # Se necessário forçar ordenação por data do arquivo:
    # import os
    # list_of_pictures.sort(key=os.path.getmtime)

    # Gera documento para cada grupo de fotos obtido.
    nomes_usados = []
    for i, grupo_de_fotos in enumerate(agrupar_fotos(fotos_da_pasta)):
        nome_do_arquivo = mostrar_fotos_obter_nome(grupo_de_fotos)

        # Se nome já tiver sido usado, gera um pós-fixo aleatório.
        if nome_do_arquivo in nomes_usados:
            nome_do_arquivo += f'_{choice(ascii_uppercase)}'  # Todo use generator
        nomes_usados.append(nome_do_arquivo)

        gerar_documento(pasta_escolhida, nome_do_arquivo, grupo_de_fotos)

        # mover fotos usadas para outra pasta.
        move_arquivos(pasta_escolhida, nome_do_arquivo, grupo_de_fotos)

        # Todo: apagar fotos pretas.


if __name__ == '__main__':
    rodar_programa()
